from __future__ import annotations
import os
from typing import Optional
from qun_alpha.models import Company, Person, Link


def _company_block_lines(c: Company) -> list[str]:
    lines = [f"### 【{c.score}·{c.status}】{c.name}  ·  {c.mentions} 次提及"]
    meta = []
    if c.sector:
        meta.append(f"赛道：{c.sector}")
    if c.stage:
        meta.append(f"阶段：{c.stage}")
    if c.financials:
        meta.append(f"融资：{c.financials}")
    if c.investors:
        meta.append(f"投资方：{', '.join(c.investors)}")
    if meta:
        lines.append("- " + " · ".join(meta))
    if c.suggested_action:
        lines.append(f"- 建议：{c.suggested_action}")
    if c.signal:
        lines.append(f"- Signal：{c.signal}")
    return lines


def _person_line(p: Person) -> str:
    bits = [p.name]
    if p.role:
        bits.append(f"（{p.role}）")
    extra = []
    if p.affiliated_companies:
        extra.append("关联：" + "、".join(p.affiliated_companies))
    if p.notable_quotes:
        extra.append("金句：" + "；".join(p.notable_quotes))
    tail = ("  ·  " + "  ·  ".join(extra)) if extra else ""
    return f"- {''.join(bits)}{tail}（{p.mentions} 次）"


def _link_line(ln: Link) -> str:
    title = ln.title or ln.url
    extra = []
    if ln.related_companies:
        extra.append("关联：" + "、".join(ln.related_companies))
    if ln.shared_by:
        extra.append("分享：" + "、".join(ln.shared_by))
    tail = ("  ·  " + "  ·  ".join(extra)) if extra else ""
    return f"- [{title}]({ln.url}){tail}"


def build_markdown(companies: list[Company], people: list[Person],
                   links: list[Link]) -> str:
    companies = sorted(companies, key=lambda c: c.score, reverse=True)
    out = ["# 群聊投资机会报告", ""]
    out.append(f"共 {len(companies)} 公司 / {len(people)} 人 / {len(links)} 链接")
    out.append("")
    out.append("## 一、公司")
    out.append("")
    if companies:
        for c in companies:
            out.extend(_company_block_lines(c))
            out.append("")
    else:
        out.append("（无）")
        out.append("")
    out.append("## 二、人物")
    out.append("")
    out.extend([_person_line(p) for p in people] or ["（无）"])
    out.append("")
    out.append("## 三、链接")
    out.append("")
    out.extend([_link_line(ln) for ln in links] or ["（无）"])
    out.append("")
    return "\n".join(out)


def _build_docx(companies: list[Company], people: list[Person],
                links: list[Link], path: str) -> None:
    from docx import Document
    companies = sorted(companies, key=lambda c: c.score, reverse=True)
    doc = Document()
    doc.add_heading("群聊投资机会报告", level=0)
    doc.add_paragraph(f"共 {len(companies)} 公司 / {len(people)} 人 / {len(links)} 链接")

    doc.add_heading("一、公司", level=1)
    for c in companies:
        doc.add_heading(f"【{c.score}·{c.status}】{c.name}  ·  {c.mentions} 次提及", level=2)
        meta = []
        if c.sector:
            meta.append(f"赛道：{c.sector}")
        if c.stage:
            meta.append(f"阶段：{c.stage}")
        if c.financials:
            meta.append(f"融资：{c.financials}")
        if c.investors:
            meta.append(f"投资方：{', '.join(c.investors)}")
        if meta:
            doc.add_paragraph(" · ".join(meta))
        if c.suggested_action:
            doc.add_paragraph(f"建议：{c.suggested_action}")
        if c.signal:
            doc.add_paragraph(f"Signal：{c.signal}")

    doc.add_heading("二、人物", level=1)
    for p in people:
        role = f"（{p.role}）" if p.role else ""
        line = f"{p.name}{role} · {p.mentions} 次"
        if p.affiliated_companies:
            line += "　关联：" + "、".join(p.affiliated_companies)
        if p.notable_quotes:
            line += "　金句：" + "；".join(p.notable_quotes)
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("三、链接", level=1)
    for ln in links:
        title = ln.title or ln.url
        line = f"{title} — {ln.url}"
        if ln.related_companies:
            line += "　关联：" + "、".join(ln.related_companies)
        if ln.shared_by:
            line += "　分享：" + "、".join(ln.shared_by)
        doc.add_paragraph(line, style="List Bullet")

    doc.save(path)


def write_reports(companies: list[Company], people: list[Person],
                  links: list[Link], out_dir: Optional[str] = None,
                  when: Optional[str] = None) -> dict:
    """生成 md + docx 两份报告，返回 {"md": path, "docx": path}。"""
    out_dir = out_dir or os.path.join(os.path.expanduser("~"), "Downloads")
    os.makedirs(out_dir, exist_ok=True)
    if when is None:
        from datetime import datetime
        when = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    md_path = os.path.join(out_dir, f"群聊投资机会_{when}.md")
    docx_path = os.path.join(out_dir, f"群聊投资机会_{when}.docx")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(build_markdown(companies, people, links))
    _build_docx(companies, people, links, docx_path)
    return {"md": md_path, "docx": docx_path}
